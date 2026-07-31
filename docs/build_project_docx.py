"""Convert a project Markdown document into a formatted Word document.

Usage:
    python -m pip install -e ".[docs]"
    python docs/build_project_docx.py SOURCE.md OUTPUT.docx
"""

from __future__ import annotations

import argparse
import re
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


HERE = Path(__file__).resolve().parent
SOURCE = HERE / "PROJECT_MASTER_GUIDE_CN.md"
OUTPUT = HERE / "PROJECT_MASTER_GUIDE_CN.docx"

BODY_FONT = "Microsoft YaHei"
CODE_FONT = "Consolas"
ACCENT = "176B55"
ACCENT_LIGHT = "E7F1ED"
INK = RGBColor(0x17, 0x21, 0x1D)
MUTED = RGBColor(0x65, 0x71, 0x6B)
INLINE_PATTERN = re.compile(r"(\*\*.+?\*\*|`[^`]+`)")


def set_run_font(run, name: str = BODY_FONT, size: float | None = None) -> None:
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = Pt(size)


def add_inline(paragraph, text: str) -> None:
    for part in INLINE_PATTERN.split(text):
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
            set_run_font(run)
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            set_run_font(run, CODE_FONT, 9)
            run.font.color.rgb = MUTED
        else:
            run = paragraph.add_run(part)
            set_run_font(run)


def strip_markdown(text: str) -> str:
    text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)
    return re.sub(r"`([^`]+)`", r"\1", text)


def shade_cell(cell, fill: str) -> None:
    properties = cell._tc.get_or_add_tcPr()
    shading = properties.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        properties.append(shading)
    shading.set(qn("w:fill"), fill)


def set_cell_text(cell, text: str, *, bold: bool = False, color=None) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(strip_markdown(text))
    run.bold = bold
    set_run_font(run, size=9)
    if color is not None:
        run.font.color.rgb = color
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run("第 ")
    set_run_font(run, size=8.5)
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = "PAGE"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend((begin, instruction, end))
    tail = paragraph.add_run(" 页")
    set_run_font(tail, size=8.5)
    for item in paragraph.runs:
        item.font.color.rgb = MUTED


def configure_document(document: Document, title: str) -> None:
    section = document.sections[0]
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.7)
    section.left_margin = Cm(1.9)
    section.right_margin = Cm(1.9)

    normal = document.styles["Normal"]
    normal.font.name = BODY_FONT
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), BODY_FONT)
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.3

    heading_sizes = {1: 16, 2: 13, 3: 11.5, 4: 10.5}
    for level, size in heading_sizes.items():
        style = document.styles[f"Heading {level}"]
        style.font.name = BODY_FONT
        style._element.rPr.rFonts.set(qn("w:eastAsia"), BODY_FONT)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor(0x17, 0x6B, 0x55)
        style.paragraph_format.space_before = Pt(12 if level <= 2 else 8)
        style.paragraph_format.space_after = Pt(5)

    document.core_properties.title = title
    document.core_properties.subject = "项目文档"
    document.core_properties.keywords = "Multi-Agent, 数据分析, 数据治理, LangGraph"

    footer = section.footer.paragraphs[0]
    add_page_number(footer)


def add_title(document: Document, text: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(36)
    paragraph.paragraph_format.space_after = Pt(18)
    run = paragraph.add_run(strip_markdown(text))
    run.bold = True
    set_run_font(run, size=22)
    run.font.color.rgb = RGBColor(0x17, 0x6B, 0x55)


def add_flow_overview(document: Document) -> None:
    caption = document.add_paragraph()
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = caption.add_run("总体业务流程")
    run.bold = True
    set_run_font(run, size=10)

    table = document.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    def merged_row(text: str, fill: str = ACCENT_LIGHT, bold: bool = True) -> None:
        cells = table.add_row().cells
        cell = cells[0].merge(cells[1])
        shade_cell(cell, fill)
        set_cell_text(cell, text, bold=bold)

    def arrow_row() -> None:
        cells = table.add_row().cells
        cell = cells[0].merge(cells[1])
        set_cell_text(cell, "↓", bold=True)

    merged_row("用户提出分析或治理目标", ACCENT, True)
    for run in table.rows[0].cells[0].paragraphs[0].runs:
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    arrow_row()
    merged_row("任务编排 Agent：理解目标、拆分任务、协调执行")
    arrow_row()
    cells = table.add_row().cells
    shade_cell(cells[0], "F4F7F5")
    shade_cell(cells[1], "F4F7F5")
    set_cell_text(cells[0], "数据分析 Agent\n只读分析、指标与案例")
    set_cell_text(cells[1], "数据治理 Agent\n契约扫描、治理问题与变更草稿")
    arrow_row()
    cells = table.add_row().cells
    set_cell_text(cells[0], "受控数据工具 / DuckDB")
    set_cell_text(cells[1], "Data Contract / Governance Adapter")
    arrow_row()
    merged_row("Observation、来源与执行轨迹返回任务编排 Agent", "F4F7F5")
    arrow_row()
    merged_row("AnswerSynthesizer → Grounding → 可信回答", ACCENT_LIGHT)
    arrow_row()
    merged_row("治理变更：独立人工审批 → Dataset Version → 验证或回滚", "FFF1DC")
    document.add_paragraph()


def split_table_row(line: str) -> list[str]:
    return [item.strip() for item in line.strip().strip("|").split("|")]


def add_table(document: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        shade_cell(cell, ACCENT)
        set_cell_text(
            cell,
            header,
            bold=True,
            color=RGBColor(0xFF, 0xFF, 0xFF),
        )
    for row_index, values in enumerate(rows):
        cells = table.add_row().cells
        for index in range(len(headers)):
            value = values[index] if index < len(values) else ""
            if row_index % 2 == 1:
                shade_cell(cells[index], "F4F7F5")
            cells[index].text = ""
            paragraph = cells[index].paragraphs[0]
            add_inline(paragraph, value)
            for run in paragraph.runs:
                set_run_font(run, size=8.5)
            cells[index].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    document.add_paragraph()


def add_code_block(document: Document, lines: list[str]) -> None:
    table = document.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    shade_cell(cell, "F4F7F5")
    cell.text = ""
    paragraph = cell.paragraphs[0]
    for index, line in enumerate(lines):
        if index:
            paragraph.add_run("\n")
        run = paragraph.add_run(line or " ")
        set_run_font(run, CODE_FONT, 8.5)
        run.font.color.rgb = MUTED
    document.add_paragraph()


def add_image(
    document: Document,
    source: Path,
    relative_path: str,
    alt_text: str,
) -> None:
    image_path = (source.parent / relative_path).resolve()
    if not image_path.exists():
        raise FileNotFoundError(f"Missing Markdown image: {image_path}")

    if image_path.name == "vehicle_quality_agent_architecture.png":
        document.add_page_break()

    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.add_run().add_picture(str(image_path), width=Cm(16.2))
    caption = document.add_paragraph()
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = caption.add_run(f"图：{alt_text}")
    run.bold = True
    run.font.color.rgb = MUTED
    set_run_font(run, size=9)

def build(source: Path = SOURCE, output: Path = OUTPUT) -> Path:
    lines = source.read_text(encoding="utf-8").splitlines()
    title = next(
        (
            strip_markdown(match.group(1))
            for line in lines
            if (match := re.match(r"^#\s+(.*)$", line.strip()))
        ),
        source.stem,
    )
    document = Document()
    configure_document(document, title)

    index = 0
    while index < len(lines):
        line = lines[index]
        stripped = line.strip()

        image = re.fullmatch(r"!\[(.*)]\((.+)\)", stripped)
        if image:
            add_image(document, source, image.group(2), image.group(1))
            index += 1
            continue

        if stripped.startswith("```"):
            language = stripped[3:].strip()
            index += 1
            block: list[str] = []
            while index < len(lines) and not lines[index].strip().startswith("```"):
                block.append(lines[index])
                index += 1
            index += 1
            if language == "mermaid":
                add_flow_overview(document)
            else:
                add_code_block(document, block)
            continue

        if (
            stripped.startswith("|")
            and index + 1 < len(lines)
            and re.fullmatch(r"[|:\- ]+", lines[index + 1].strip())
        ):
            headers = split_table_row(stripped)
            index += 2
            rows: list[list[str]] = []
            while index < len(lines) and lines[index].strip().startswith("|"):
                rows.append(split_table_row(lines[index]))
                index += 1
            add_table(document, headers, rows)
            continue

        heading = re.match(r"^(#{1,6})\s+(.*)$", stripped)
        if heading:
            level = len(heading.group(1))
            text = heading.group(2)
            if level == 1:
                add_title(document, text)
            else:
                document.add_heading(strip_markdown(text), level=min(level - 1, 4))
            index += 1
            continue

        if stripped in {"---", "***", "___"}:
            index += 1
            continue

        if stripped.startswith(">"):
            paragraph = document.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.left_indent = Cm(1)
            paragraph.paragraph_format.right_indent = Cm(1)
            add_inline(paragraph, stripped.lstrip("> "))
            for run in paragraph.runs:
                run.italic = True
                run.font.color.rgb = MUTED
            index += 1
            continue

        bullet = re.match(r"^\s*[-*]\s+(.*)$", line)
        numbered = re.match(r"^\s*\d+\.\s+(.*)$", line)
        if bullet or numbered:
            match = bullet or numbered
            style = "List Bullet" if bullet else "List Number"
            paragraph = document.add_paragraph(style=style)
            add_inline(paragraph, match.group(1))
            index += 1
            continue

        if not stripped:
            index += 1
            continue

        paragraph = document.add_paragraph()
        add_inline(paragraph, stripped)
        index += 1

    document.save(output)
    return output


if __name__ == "__main__":
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("source", nargs="?", type=Path, default=SOURCE)
    parser.add_argument("output", nargs="?", type=Path, default=OUTPUT)
    arguments = parser.parse_args()
    print(f"wrote {build(arguments.source.resolve(), arguments.output.resolve())}")